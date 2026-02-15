#!/usr/bin/env python3
"""
build_resume_onepage.py (production)

Markdown resume -> ATS-safe DOCX -> PDF (LibreOffice), enforcing 1 page.

Polish:
- Clickable links (mailto/tel/https)
- Clean contact line: Email | Phone | LinkedIn | GitHub
- H3 headers bold (no literal **)
- Ignore --- horizontal rules
- Bullets with hanging indent + tight spacing
- Trim rules resilient to **bold** in H3 titles

Usage:
  python build_resume_onepage.py resume_helpdesk.md out.docx out.pdf

Requires:
  pip install python-docx pypdf
  LibreOffice installed (soffice.exe)
"""

import re
import sys
import shutil
import subprocess
from pathlib import Path
from typing import List, Tuple, Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pypdf import PdfReader


# ----------------------------
# Trimming rules (only used if PDF > 1 page)
# ----------------------------
TRIM_RULES = [
    ("DROP_SECTION", "Other Professional Experience"),
    ("KEEP_ONLY_PROJECTS", 2),
    ("TRIM_BULLETS_UNDER", ("Technical & IT Support Roles", 3)),
    ("TRIM_BULLETS_UNDER", ("Windows Event Monitoring & Mini SOC Lab", 2)),
    ("TRIM_BULLETS_UNDER", ("pfSense Firewall & Network Segmentation Lab", 2)),
    ("TRIM_BULLETS_UNDER", ("Small Business IT & Security Assessments", 1)),
]

# ----------------------------
# Formatting presets (tries tighter if needed)
# ----------------------------
FMT_PRESETS = [
    dict(
        margin=0.50,
        body_pt=10.0,
        name_pt=13.0,
        h2_pt=10.6,
        h3_pt=10.2,
        h2_before=6,
        h2_after=2,
        h3_before=4,
        h3_after=1,
        para_after=1,
        bullet_after=0,
        bullet_left=0.25,   # inches
        bullet_hang=0.18,   # inches
    ),
    dict(
        margin=0.45,
        body_pt=9.8,
        name_pt=12.5,
        h2_pt=10.4,
        h3_pt=10.0,
        h2_before=5,
        h2_after=2,
        h3_before=3,
        h3_after=1,
        para_after=0,
        bullet_after=0,
        bullet_left=0.25,
        bullet_hang=0.18,
    ),
]

# ----------------------------
# Regex helpers
# ----------------------------
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
H2_RE = re.compile(r"^##\s+(.+)$")
H3_RE = re.compile(r"^###\s+(.+)$")
BULLET_RE = re.compile(r"^\s*-\s+(.+)$")
HR_RE = re.compile(r"^\s*-{3,}\s*$")

EMAIL_RE = re.compile(r"(?P<email>[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,})")
PHONE_RE = re.compile(
    r"(?P<phone>(?:\+?1[\s\-\.]?)?(?:\(\d{3}\)|\d{3})[\s\-\.]?\d{3}[\s\-\.]?\d{4})"
)
URL_RE = re.compile(
    r"(?P<url>(?:https?://)?(?:www\.)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s|,)]+)?)"
)


def strip_md_bold(s: str) -> str:
    return re.sub(r"\*\*(.+?)\*\*", r"\1", s).strip()


def normalize_url(url: str) -> str:
    u = url.strip()
    if u.startswith("http://") or u.startswith("https://"):
        return u
    return "https://" + u


def normalize_tel(phone: str) -> str:
    # Best effort E.164-ish for tel: links
    digits = re.sub(r"\D", "", phone)
    if len(digits) == 10:
        digits = "1" + digits
    if digits.startswith("1") and len(digits) == 11:
        return f"+{digits}"
    return digits


def find_soffice() -> str:
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.com",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.com",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    lo = shutil.which("soffice") or shutil.which("soffice.exe") or shutil.which("soffice.com")
    if lo:
        return lo
    raise RuntimeError("LibreOffice not found. Install LibreOffice or add soffice to PATH.")


def pdf_pages(pdf_path: Path) -> int:
    return len(PdfReader(str(pdf_path)).pages)


def md_read(md_path: Path) -> List[str]:
    return md_path.read_text(encoding="utf-8").splitlines()


def normalize_whitespace(lines: List[str]) -> List[str]:
    cleaned: List[str] = []
    blank = 0
    for ln in lines:
        ln = re.sub(r"\s{2,}$", "", ln.rstrip())  # remove forced MD linebreak spaces
        if HR_RE.match(ln):
            continue
        if ln.strip() == "":
            blank += 1
            if blank <= 1:
                cleaned.append("")
        else:
            blank = 0
            cleaned.append(ln)
    return cleaned


# ----------------------------
# Trim logic
# ----------------------------
def drop_section(lines: List[str], section_title: str) -> List[str]:
    out: List[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        m3 = H3_RE.match(line)
        if m3 and strip_md_bold(m3.group(1)) == section_title:
            i += 1
            while i < len(lines) and not H3_RE.match(lines[i]) and not H2_RE.match(lines[i]):
                i += 1
            continue
        out.append(line)
        i += 1
    return out


def keep_only_projects(lines: List[str], keep_n: int) -> List[str]:
    out: List[str] = []
    i = 0
    in_projects = False
    project_count = 0
    skipping = False

    while i < len(lines):
        line = lines[i]
        if H2_RE.match(line) and "TECHNICAL PROJECT" in line.upper():
            in_projects = True
            skipping = False
            out.append(line)
            i += 1
            continue

        if in_projects and H2_RE.match(line):
            in_projects = False
            skipping = False
            out.append(line)
            i += 1
            continue

        if in_projects and H3_RE.match(line):
            project_count += 1
            skipping = project_count > keep_n
            if not skipping:
                out.append(line)
            i += 1
            continue

        if in_projects and skipping:
            i += 1
            continue

        out.append(line)
        i += 1

    return out


def trim_bullets_under(lines: List[str], header_title: str, keep_k: int) -> List[str]:
    out: List[str] = []
    i = 0
    in_target = False
    kept = 0

    while i < len(lines):
        line = lines[i]
        if H3_RE.match(line):
            title = strip_md_bold(H3_RE.match(line).group(1))
            in_target = (title == header_title)
            kept = 0
            out.append(line)
            i += 1
            continue

        if in_target:
            bm = BULLET_RE.match(line)
            if bm:
                kept += 1
                if kept <= keep_k:
                    out.append(line)
                i += 1
                continue
            out.append(line)
            i += 1
            continue

        out.append(line)
        i += 1

    return out


def apply_trim_rule(lines: List[str], rule: Tuple) -> List[str]:
    kind = rule[0]
    if kind == "DROP_SECTION":
        return drop_section(lines, rule[1])
    if kind == "KEEP_ONLY_PROJECTS":
        return keep_only_projects(lines, rule[1])
    if kind == "TRIM_BULLETS_UNDER":
        title, keep_k = rule[1]
        return trim_bullets_under(lines, title, keep_k)
    return lines


# ----------------------------
# Word hyperlink helpers
# ----------------------------
def add_hyperlink(paragraph, url: str, text: str, *, bold: bool = False, font_size: Optional[Pt] = None):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if bold:
        rPr.append(OxmlElement("w:b"))

    if font_size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(font_size.pt * 2)))
        rPr.append(sz)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_runs_with_bold(p, text: str, *, force_bold: bool = False, font_size: Optional[Pt] = None):
    """
    Adds text supporting **bold** and hyperlinks (email/phone/url).
    Email > phone > url precedence for overlapping matches.
    """
    # Split by **bold**
    segments = []
    last = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > last:
            segments.append((False, text[last:m.start()]))
        segments.append((True, m.group(1)))
        last = m.end()
    if last < len(text):
        segments.append((False, text[last:]))

    def add_plain(s: str, is_bold: bool):
        if not s:
            return
        run = p.add_run(s)
        run.bold = force_bold or is_bold
        if font_size is not None:
            run.font.size = font_size

    for is_bold_seg, seg in segments:
        seg = seg or ""
        i = 0
        while i < len(seg):
            em = EMAIL_RE.search(seg, i)
            pm = PHONE_RE.search(seg, i)
            um = URL_RE.search(seg, i)

            # Choose earliest; tie-break: email > phone > url
            candidates = []
            if em: candidates.append(("email", em.start(), em))
            if pm: candidates.append(("phone", pm.start(), pm))
            if um: candidates.append(("url", um.start(), um))
            if not candidates:
                add_plain(seg[i:], is_bold_seg)
                break

            candidates.sort(key=lambda x: (x[1], {"email": 0, "phone": 1, "url": 2}[x[0]]))
            kind, start, m = candidates[0]

            if start > i:
                add_plain(seg[i:start], is_bold_seg)

            token = m.group(0)
            bold = force_bold or is_bold_seg

            if kind == "email":
                email = m.group("email")
                add_hyperlink(p, f"mailto:{email}", email, bold=bold, font_size=font_size)
            elif kind == "phone":
                phone = m.group("phone")
                tel = normalize_tel(phone)
                # Keep visible formatting; link uses tel: normalized
                add_hyperlink(p, f"tel:{tel}", phone, bold=bold, font_size=font_size)
            else:
                raw_url = m.group("url")
                if "@" in raw_url:
                    add_plain(raw_url, is_bold_seg)
                else:
                    add_hyperlink(p, normalize_url(raw_url), raw_url, bold=bold, font_size=font_size)

            i = m.end()


# ----------------------------
# Contact line normalization
# ----------------------------
def is_contact_block(lines: List[str], idx: int) -> bool:
    # We expect after "# NAME" a few lines with location/contact.
    # We'll normalize lines until we hit "## " or blank line after contact.
    if idx < 0 or idx >= len(lines):
        return False
    return True


def normalize_contact_lines(lines: List[str]) -> List[str]:
    """
    If the top area includes email/phone/linkedin/github, render a clean clickable contact line:
      Email | Phone | LinkedIn | GitHub
    This is done in DOCX creation stage by recognizing 'contact mode' lines.
    Here we only return lines unchanged; the DOCX writer has special handling.
    """
    return lines


# ----------------------------
# DOCX creation
# ----------------------------
def md_to_docx(lines: List[str], docx_path: Path, fmt: dict):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(fmt["margin"])
    sec.bottom_margin = Inches(fmt["margin"])
    sec.left_margin = Inches(fmt["margin"])
    sec.right_margin = Inches(fmt["margin"])

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(fmt["body_pt"])

    # Precompute bullet indent values
    bullet_left = Inches(fmt["bullet_left"])
    bullet_hang = Inches(fmt["bullet_hang"])

    # Contact normalization state:
    in_header_block = False
    header_lines: List[str] = []

    def flush_header_block():
        nonlocal header_lines, in_header_block
        if not header_lines:
            return

        # Build one or two header lines:
        # - location line stays as plain text
        # - contact line becomes Email | Phone | LinkedIn | GitHub (clickable)
        # We’ll parse tokens across header_lines.
        text = " ".join(header_lines)
        # Extract
        email = EMAIL_RE.search(text)
        phone = PHONE_RE.search(text)
        urls = URL_RE.findall(text)

        # Identify LinkedIn/GitHub (best effort)
        linkedin = None
        github = None
        other_urls = []

        for u in urls:
            u_norm = u.lower()
            if "linkedin.com" in u_norm and linkedin is None:
                linkedin = u
            elif "github.com" in u_norm and github is None:
                github = u
            else:
                other_urls.append(u)

        # First line: try to keep a location line if present
        # (heuristic: first header line often is "Omaha, NE")
        # We'll just print the first header line that doesn't contain email/phone/url.
        location_line = None
        for hl in header_lines:
            if (EMAIL_RE.search(hl) is None) and (PHONE_RE.search(hl) is None) and (URL_RE.search(hl) is None):
                location_line = hl.strip()
                break

        if location_line:
            p = doc.add_paragraph()
            add_runs_with_bold(p, location_line, force_bold=False)
            p.paragraph_format.space_after = Pt(0)

        # Contact line: Email | Phone | LinkedIn | GitHub
        parts = []
        if email:
            parts.append(("email", email.group("email")))
        if phone:
            parts.append(("phone", phone.group("phone")))
        if linkedin:
            parts.append(("url", linkedin))
        if github:
            parts.append(("url", github))

        # If nothing extracted, just output original lines
        if not parts:
            for hl in header_lines:
                p = doc.add_paragraph()
                add_runs_with_bold(p, hl.strip(), force_bold=False)
                p.paragraph_format.space_after = Pt(0)
        else:
            p = doc.add_paragraph()
            # Build clickable chunks with separators
            first = True
            for kind, val in parts:
                if not first:
                    run = p.add_run(" | ")
                    run.bold = False
                first = False

                if kind == "email":
                    add_hyperlink(p, f"mailto:{val}", val, bold=False, font_size=None)
                elif kind == "phone":
                    add_hyperlink(p, f"tel:{normalize_tel(val)}", val, bold=False, font_size=None)
                else:
                    label = "LinkedIn" if linkedin and val == linkedin else ("GitHub" if github and val == github else val)
                    add_hyperlink(p, normalize_url(val), label, bold=False, font_size=None)

            p.paragraph_format.space_after = Pt(2)

        header_lines = []
        in_header_block = False

    for i, line in enumerate(lines):
        line = line.rstrip()
        if not line.strip():
            # if we're in header block, stop it on blank line
            if in_header_block:
                flush_header_block()
            continue
        if HR_RE.match(line):
            continue

        # Name line starts header parsing
        if line.startswith("# "):
            # flush any previous header (shouldn’t happen)
            if in_header_block:
                flush_header_block()

            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(fmt["name_pt"])
            p.paragraph_format.space_after = Pt(1)

            in_header_block = True
            header_lines = []
            continue

        # Capture header block lines until first ## section
        if in_header_block:
            if line.startswith("## "):
                flush_header_block()
                # fall through to section rendering
            else:
                header_lines.append(line.strip())
                continue

        # Sections
        if line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(fmt["h2_pt"])
            p.paragraph_format.space_before = Pt(fmt["h2_before"])
            p.paragraph_format.space_after = Pt(fmt["h2_after"])
            continue

        # Subheaders (roles/projects) - bold always
        if line.startswith("### "):
            p = doc.add_paragraph()
            title = line[4:].strip()
            add_runs_with_bold(p, title, force_bold=True, font_size=Pt(fmt["h3_pt"]))
            p.paragraph_format.space_before = Pt(fmt["h3_before"])
            p.paragraph_format.space_after = Pt(fmt["h3_after"])
            continue

        # Bullets with hanging indent
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, line[2:].strip(), force_bold=False, font_size=None)
            pf = p.paragraph_format
            pf.left_indent = bullet_left
            pf.first_line_indent = -bullet_hang
            pf.space_after = Pt(fmt["bullet_after"])
            continue

        # Normal paragraphs (links enabled)
        p = doc.add_paragraph()
        add_runs_with_bold(p, line.strip(), force_bold=False, font_size=None)
        p.paragraph_format.space_after = Pt(fmt["para_after"])

    # flush header if file ends before a section
    if in_header_block:
        flush_header_block()

    doc.save(str(docx_path))


def docx_to_pdf(docx_path: Path, pdf_path: Path):
    soffice = find_soffice()
    outdir = pdf_path.parent
    outdir.mkdir(parents=True, exist_ok=True)

    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nofirststartwizard",
        "--convert-to", "pdf",
        "--outdir", str(outdir),
        str(docx_path),
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

    generated = outdir / (docx_path.stem + ".pdf")
    if not generated.exists():
        raise RuntimeError("LibreOffice conversion ran, but PDF was not created.")
    if generated.resolve() != pdf_path.resolve():
        generated.replace(pdf_path)


def main():
    if len(sys.argv) != 4:
        print("Usage: python build_resume_onepage.py resume.md out.docx out.pdf")
        sys.exit(2)

    md_path = Path(sys.argv[1]).expanduser().resolve()
    docx_path = Path(sys.argv[2]).expanduser().resolve()
    pdf_path = Path(sys.argv[3]).expanduser().resolve()

    if not md_path.exists():
        print(f"Markdown not found: {md_path}")
        sys.exit(1)

    base_lines = normalize_whitespace(md_read(md_path))
    attempts = 0

    for fmt in FMT_PRESETS:
        cur = list(base_lines)

        # Try without trimming first
        md_to_docx(cur, docx_path, fmt)
        docx_to_pdf(docx_path, pdf_path)
        pages = pdf_pages(pdf_path)
        attempts += 1

        if pages == 1:
            print(f"OK: 1 page (no trimming). Attempts: {attempts}")
            return

        # Apply trim rules until 1 page
        for rule in TRIM_RULES:
            cur = normalize_whitespace(apply_trim_rule(cur, rule))

            md_to_docx(cur, docx_path, fmt)
            docx_to_pdf(docx_path, pdf_path)
            pages = pdf_pages(pdf_path)
            attempts += 1

            if pages == 1:
                print(f"OK: 1 page after trimming rule: {rule}. Attempts: {attempts}")
                return

    print("Could not reach 1 page with current rules.")
    print("Last output was generated anyway; consider trimming one more bullet under IT Support Roles.")
    sys.exit(3)


if __name__ == "__main__":
    main()
